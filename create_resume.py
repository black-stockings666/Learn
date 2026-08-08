from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from PIL import Image
from pathlib import Path

ROOT = Path(r"S:\videonest")
OUT = ROOT / "李飞洋_Java后端开发实习生简历.docx"
SOURCE = Path(r"C:\Users\我爱cb\AppData\Local\Temp\codex-clipboard-dbc3bb24-2ae7-49c3-9c0f-1f8bc4409ab9.jpg")
PHOTO = ROOT / "resume_photo.png"


def set_run_font(run, name="Microsoft YaHei", size=10.5, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")


def set_table_fixed(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblLayout = tblPr.first_child_found_in("w:tblLayout")
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)
    tblLayout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def add_bottom_border(paragraph, color="222222", size="8"):
    pPr = paragraph._p.get_or_add_pPr()
    borders = pPr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        pPr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_section(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(title)
    set_run_font(r, size=13.5, bold=True, color=(0, 0, 0))
    add_bottom_border(p)
    return p


def add_line(doc, text, size=9.6, after=2.5, color=(0, 0, 0), indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.16
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    set_run_font(r, size=size, color=color)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1.8)
    p.paragraph_format.line_spacing = 1.13
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.first_line_indent = Cm(-0.45)
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    set_run_font(r, size=9.35)
    return p


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rPr.append(rFonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2F64E8")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def main():
    # Named override: A4 portrait and 1.45 cm margins for a one-page Chinese resume.
    image = Image.open(SOURCE)
    # Center crop to the resume photo ratio (about 2:3) without stretching.
    image.crop((26, 0, 454, 640)).save(PHOTO)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.1)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.6)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.16

    # Header block: fixed two-column geometry avoids photo overlap.
    header = doc.add_table(rows=1, cols=2)
    set_table_fixed(header, [7600, 1240])
    left, right = header.rows[0].cells
    set_cell_margins(left, 0, 0, 0, 0)
    set_cell_margins(right, 0, 0, 0, 0)
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("李飞洋")
    set_run_font(r, size=20, bold=True)
    for text in [
        "电话：19560707895    邮箱：1024794515@gmail.com",
        "求职意向：Java 后端实习生",
    ]:
        p = left.add_paragraph()
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(text)
        set_run_font(r, size=10.1)
    p = left.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("GitHub：")
    set_run_font(r, size=10.1)
    add_hyperlink(p, "https://github.com/black-stockings666", "https://github.com/black-stockings666")
    pic = right.paragraphs[0]
    pic.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pic.paragraph_format.space_after = Pt(0)
    pic.add_run().add_picture(str(PHOTO), width=Cm(1.78), height=Cm(2.66))

    add_section(doc, "教育经历")
    edu = doc.add_table(rows=1, cols=2)
    set_table_fixed(edu, [6750, 2090])
    a, b = edu.rows[0].cells
    set_cell_margins(a, 0, 0, 0, 0)
    set_cell_margins(b, 0, 0, 0, 0)
    p = a.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("山东科技大学")
    set_run_font(r, size=11.4, bold=True)
    r = p.add_run("    计算机科学与技术    本科")
    set_run_font(r, size=10, color=(105, 105, 105))
    p = a.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("全日制    山东省青岛市黄岛区")
    set_run_font(r, size=9.6, color=(105, 105, 105))
    p = b.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("2024-09 ~ 至今")
    set_run_font(r, size=10, color=(105, 105, 105))

    add_section(doc, "专业技能")
    skills = [
        "熟悉 Java 基础、面向对象思想、集合框架、异常处理及常用设计模式。",
        "熟悉 Spring Boot、Spring MVC、Spring Security，能够完成 RESTful API 设计与 JWT 鉴权。",
        "熟悉 MySQL 索引设计、事务与 SQL 优化，使用 MyBatis-Plus 完成 CRUD 和复杂查询开发。",
        "熟悉 Redis、Redis ZSet 与 Lua 脚本，实践热榜、计数、去重、限流及缓存问题治理。",
        "熟悉 RabbitMQ、MinIO 对象存储和预签名 URL 上传；了解 FFmpeg 视频转码与异步任务编排。",
        "熟练使用 IntelliJ IDEA、VS Code、Git、Maven、Postman；具备使用 Codex 辅助分析、编码、调试与重构的实践经验。",
    ]
    for item in skills:
        add_numbered(doc, item)

    add_section(doc, "项目经历")
    project = doc.add_table(rows=1, cols=2)
    set_table_fixed(project, [6000, 2840])
    a, b = project.rows[0].cells
    set_cell_margins(a, 0, 0, 0, 0)
    set_cell_margins(b, 0, 0, 0, 0)
    p = a.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("VideoNest 视频社区平台")
    set_run_font(r, size=11.4, bold=True)
    p = a.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("个人项目｜全栈开发，后端为主")
    set_run_font(r, size=9.4, color=(105, 105, 105))
    p = b.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("2026-07 ~ 至今")
    set_run_font(r, size=10, color=(105, 105, 105))
    p = b.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    add_hyperlink(p, "GitHub 项目地址", "https://github.com/black-stockings666")

    bullets = [
        "技术栈：Java、Spring Boot、Spring Security、MyBatis-Plus、MySQL、Redis、RabbitMQ、MinIO、FFmpeg、Docker Compose、Nginx。",
        "独立开发视频社区核心业务，基于 JWT + RBAC 实现登录鉴权与权限控制，完成视频投稿、审核发布、点赞收藏、评论回复、关注和消息通知等 RESTful 接口。",
        "设计视频上传处理链路：通过 MinIO 预签名 URL 实现前端直传，使用 RabbitMQ 异步处理 FFmpeg 转码、封面生成与审核超时任务，降低接口等待时间。",
        "基于 Redis + Lua 脚本实现播放量去重、匿名限流与热度统计；采用小时分桶与 6 小时半衰期计算热门视频榜单，并配置降级策略。",
        "通过随机 TTL、空值缓存和互斥锁治理缓存穿透、击穿与雪崩；设计发布流、评论分页等复合索引并消除评论列表 N+1 查询。",
        "完善消息队列可靠性机制，配置消费重试、延迟消息和死信队列；使用 Docker Compose 编排全套服务，实现环境快速部署与复现。",
    ]
    for item in bullets:
        add_numbered(doc, item)

    # Quiet footer, deliberately minimal for a one-page resume.
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    r = fp.add_run("Java 后端开发实习生简历")
    set_run_font(r, size=7.5, color=(140, 140, 140))

    doc.core_properties.author = "李飞洋"
    doc.core_properties.title = "Java 后端开发实习生简历"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
